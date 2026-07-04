# metarag/core/loader.py

import os
from typing import List
from pathlib import Path
from abc import ABC, abstractmethod


class Document:
    """Simple document representation."""
    
    def __init__(self, text: str, metadata: dict = None):
        self.text = text
        self.metadata = metadata or {}
    
    def __repr__(self):
        return f"Document({len(self.text)} chars, {self.metadata})"


class LoaderInterface(ABC):
    """Contract for document loaders."""
    
    @abstractmethod
    def load(self) -> List[Document]:
        """Load documents from source."""
        pass


class DocumentLoader(LoaderInterface):
    """
    Load documents from various formats.
    No hard dependencies — tries to import only when needed.
    Guides user if format not supported.
    """
    
    def __init__(self, path: str):
        self.path = Path(path)
        if not self.path.exists():
            raise FileNotFoundError(f"Path not found: {path}")
    
    def load(self) -> List[Document]:
        """Load all documents from path (directory or file)."""
        docs = []
        
        if self.path.is_file():
            docs.extend(self._load_file(self.path))
        elif self.path.is_dir():
            for file_path in self.path.rglob("*"):
                if file_path.is_file():
                    docs.extend(self._load_file(file_path))
        
        if not docs:
            print(f"[DocumentLoader] No documents loaded from {self.path}")
        else:
            print(f"[DocumentLoader] Loaded {len(docs)} documents")
        
        return docs
    
    def _load_file(self, file_path: Path) -> List[Document]:
        """Load single file based on extension."""
        ext = file_path.suffix.lower()
        
        try:
            if ext == ".pdf":
                return self._load_pdf(file_path)
            elif ext == ".txt":
                return self._load_txt(file_path)
            elif ext == ".docx":
                return self._load_docx(file_path)
            elif ext == ".html" or ext == ".htm":
                return self._load_html(file_path)
            elif ext == ".json":
                return self._load_json(file_path)
            elif ext == ".csv":
                return self._load_csv(file_path)
            elif ext == ".md":
                return self._load_markdown(file_path)
            else:
                # Silent skip for unsupported formats
                return []
        except Exception as e:
            print(f"[DocumentLoader] Error loading {file_path.name}: {e}")
            return []
    
    # ─────────────────────────────────────────────────────────
    # PDF Loading
    # ─────────────────────────────────────────────────────────
    
    def _load_pdf(self, file_path: Path) -> List[Document]:
        """Load PDF (optional: pypdf)."""
        try:
            import pypdf
        except ImportError:
            print(f"[DocumentLoader] ⚠️  pypdf required for PDFs. Install: pip install pypdf")
            return []
        
        docs = []
        try:
            reader = pypdf.PdfReader(file_path)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    docs.append(Document(
                        text=text,
                        metadata={
                            "source": file_path.name,
                            "page": i,
                            "type": "pdf"
                        }
                    ))
        except Exception as e:
            print(f"[DocumentLoader] Error parsing PDF {file_path.name}: {e}")
        
        return docs
    
    # ─────────────────────────────────────────────────────────
    # Text Files
    # ─────────────────────────────────────────────────────────
    
    def _load_txt(self, file_path: Path) -> List[Document]:
        """Load plain text."""
        try:
            text = file_path.read_text(encoding="utf-8")
            if text.strip():
                return [Document(
                    text=text,
                    metadata={"source": file_path.name, "type": "txt"}
                )]
        except UnicodeDecodeError:
            print(f"[DocumentLoader] Could not decode {file_path.name} (not UTF-8)")
        
        return []
    
    # ─────────────────────────────────────────────────────────
    # DOCX Files
    # ─────────────────────────────────────────────────────────
    
    def _load_docx(self, file_path: Path) -> List[Document]:
        """Load DOCX (optional: python-docx)."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            print(f"[DocumentLoader] ⚠️  python-docx required for DOCX. Install: pip install python-docx")
            return []
        
        try:
            doc = DocxDocument(file_path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            
            if text.strip():
                return [Document(
                    text=text,
                    metadata={"source": file_path.name, "type": "docx"}
                )]
        except Exception as e:
            print(f"[DocumentLoader] Error parsing DOCX {file_path.name}: {e}")
        
        return []
    
    # ─────────────────────────────────────────────────────────
    # HTML Files
    # ─────────────────────────────────────────────────────────
    
    def _load_html(self, file_path: Path) -> List[Document]:
        """Load HTML (optional: beautifulsoup4)."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            print(f"[DocumentLoader] ⚠️  beautifulsoup4 required for HTML. Install: pip install beautifulsoup4")
            return []
        
        try:
            html = file_path.read_text(encoding="utf-8")
            soup = BeautifulSoup(html, "html.parser")
            
            # Remove script and style elements
            for tag in soup(["script", "style", "nav", "footer"]):
                tag.decompose()
            
            text = soup.get_text(separator="\n", strip=True)
            text = "\n".join(line.strip() for line in text.split("\n") if line.strip())
            
            if text.strip():
                return [Document(
                    text=text,
                    metadata={"source": file_path.name, "type": "html"}
                )]
        except Exception as e:
            print(f"[DocumentLoader] Error parsing HTML {file_path.name}: {e}")
        
        return []
    
    # ─────────────────────────────────────────────────────────
    # JSON Files
    # ─────────────────────────────────────────────────────────
    
    def _load_json(self, file_path: Path) -> List[Document]:
        """Load JSON (assumes list of dicts with 'text' field)."""
        import json
        
        try:
            data = json.loads(file_path.read_text())
            docs = []
            
            if isinstance(data, list):
                for i, item in enumerate(data):
                    if isinstance(item, dict) and "text" in item:
                        metadata = item.get("metadata", {})
                        metadata["source"] = file_path.name
                        metadata["type"] = "json"
                        metadata["index"] = i
                        
                        docs.append(Document(text=item["text"], metadata=metadata))
            
            return docs
        except json.JSONDecodeError as e:
            print(f"[DocumentLoader] Invalid JSON in {file_path.name}: {e}")
        except Exception as e:
            print(f"[DocumentLoader] Error parsing JSON {file_path.name}: {e}")
        
        return []
    
    # ─────────────────────────────────────────────────────────
    # CSV Files
    # ─────────────────────────────────────────────────────────
    
    def _load_csv(self, file_path: Path) -> List[Document]:
        """Load CSV (optional: pandas, assumes 'text' column)."""
        try:
            import pandas as pd
        except ImportError:
            print(f"[DocumentLoader] ⚠️  pandas required for CSV. Install: pip install pandas")
            return []
        
        try:
            df = pd.read_csv(file_path)
            
            if "text" not in df.columns:
                print(f"[DocumentLoader] CSV must have 'text' column. Found: {list(df.columns)}")
                return []
            
            docs = []
            for idx, row in df.iterrows():
                metadata = {col: str(row[col]) for col in df.columns if col != "text"}
                metadata["source"] = file_path.name
                metadata["type"] = "csv"
                metadata["row"] = idx
                
                docs.append(Document(text=row["text"], metadata=metadata))
            
            return docs
        except Exception as e:
            print(f"[DocumentLoader] Error parsing CSV {file_path.name}: {e}")
        
        return []
    
    # ─────────────────────────────────────────────────────────
    # Markdown Files
    # ─────────────────────────────────────────────────────────
    
    def _load_markdown(self, file_path: Path) -> List[Document]:
        """Load Markdown."""
        try:
            text = file_path.read_text(encoding="utf-8")
            if text.strip():
                return [Document(
                    text=text,
                    metadata={"source": file_path.name, "type": "markdown"}
                )]
        except Exception as e:
            print(f"[DocumentLoader] Error parsing Markdown {file_path.name}: {e}")
        
        return []
    
    # ─────────────────────────────────────────────────────────
    # Load by Format
    # ─────────────────────────────────────────────────────────
    
    def load_pdfs(self) -> List[Document]:
        """Load only PDFs."""
        return [doc for doc in self.load() if doc.metadata.get("type") == "pdf"]
    
    def load_texts(self) -> List[Document]:
        """Load only text files."""
        return [doc for doc in self.load() if doc.metadata.get("type") == "txt"]
    
    def load_jsons(self) -> List[Document]:
        """Load only JSON files."""
        return [doc for doc in self.load() if doc.metadata.get("type") == "json"]
    
    def load_format(self, fmt: str) -> List[Document]:
        """Load specific format."""
        return [doc for doc in self.load() if doc.metadata.get("type") == fmt.lower()]